#!/usr/bin/env python3
"""Build readable DigiStayBook BOP artifacts from the Markdown authority.

Outputs:
- a styled DOCX using the compact_reference_guide preset;
- a standalone readable HTML view of v3;
- a standalone v2-to-v3 HTML change review.
"""

from __future__ import annotations
import html

import argparse
import difflib
import math
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


REQ_RE = re.compile(r"\s+\[(DSB-BOP-P\d+-[0-9A-Z]+)\]\s*$")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
TABLE_SEPARATOR_RE = re.compile(r"^\|(?:\s*:?-+:?\s*\|)+\s*$")
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|(?<!\*)\*[^*]+?\*(?!\*))")

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 100, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"


@dataclass
class Block:
    kind: str
    value: object
    level: int = 0


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def parse_markdown(text: str) -> list[Block]:
    lines = text.splitlines()
    blocks: list[Block] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        heading = HEADING_RE.match(line)
        if heading:
            blocks.append(Block("heading", heading.group(2), len(heading.group(1))))
            i += 1
            continue
        if stripped == "---":
            blocks.append(Block("rule", ""))
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and TABLE_SEPARATOR_RE.match(lines[i + 1]):
            rows = [parse_table_row(line)]
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            blocks.append(Block("table", rows))
            continue
        if re.match(r"^\s*-\s+", line):
            items: list[str] = []
            while i < len(lines) and re.match(r"^\s*-\s+", lines[i]):
                items.append(re.sub(r"^\s*-\s+", "", lines[i]).strip())
                i += 1
            blocks.append(Block("bullets", items))
            continue
        if re.match(r"^\s*\d+\.\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                items.append(re.sub(r"^\s*\d+\.\s+", "", lines[i]).strip())
                i += 1
            blocks.append(Block("numbers", items))
            continue
        if line.startswith(">"):
            quote: list[str] = []
            while i < len(lines) and lines[i].startswith(">"):
                quote.append(lines[i][1:].strip())
                i += 1
            blocks.append(Block("quote", " ".join(quote)))
            continue

        if line.startswith("`"):
            lang = line[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("`"):
                code_lines.append(lines[i].rstrip())
                i += 1
            i += 1 # skip closing `
            blocks.append(Block("code", {"lang": lang, "code": "\n".join(code_lines)}))
            continue

        paragraph = [stripped]
        i += 1
        while i < len(lines):
            candidate = lines[i]
            if not candidate.strip():
                break
            if (
                HEADING_RE.match(candidate)
                or candidate.strip() == "---"
                or candidate.startswith("|")
                or candidate.startswith(">")
                or re.match(r"^\s*-\s+", candidate)
                or re.match(r"^\s*\d+\.\s+", candidate)
            ):
                break
            paragraph.append(candidate.strip())
            i += 1
        blocks.append(Block("paragraph", " ".join(paragraph)))
    return blocks


def split_requirement(text: str) -> tuple[str, str | None]:
    match = REQ_RE.search(text)
    if not match:
        return text, None
    return text[: match.start()].rstrip(), match.group(1)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_inline_runs(paragraph, text: str, *, size: float | None = None):
    position = 0
    for token in INLINE_RE.finditer(text):
        if token.start() > position:
            run = paragraph.add_run(text[position : token.start()])
            set_run_font(run, size=size)
        raw = token.group(0)
        if raw.startswith("**"):
            run = paragraph.add_run(raw[2:-2])
            set_run_font(run, size=size, bold=True)
        elif raw.startswith("`"):
            run = paragraph.add_run(raw[1:-1])
            set_run_font(run, name="Consolas", size=(size or 11) - 0.5, color=DARK_BLUE)
        else:
            run = paragraph.add_run(raw[1:-1])
            set_run_font(run, size=size, italic=True)
        position = token.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size)


def set_cell_fill(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_width(parent, tag: str, width: int):
    node = ensure_child(parent, tag)
    node.set(qn("w:type"), "dxa")
    node.set(qn("w:w"), str(width))


def apply_table_geometry(table, widths: list[int], total_width: int, indent: int, margins: dict[str, int]):
    if sum(widths) != total_width:
        raise ValueError("Table column widths must equal the table width")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    set_width(tbl_pr, "w:tblW", total_width)
    tbl_indent = ensure_child(tbl_pr, "w:tblInd")
    tbl_indent.set(qn("w:type"), "dxa")
    tbl_indent.set(qn("w:w"), str(indent))
    ensure_child(tbl_pr, "w:tblLayout").set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for column_index, width in enumerate(widths):
        table.columns[column_index].width = Twips(width)
    for row in table.rows:
        row.height = None
        for column_index, cell in enumerate(row.cells):
            width = widths[column_index]
            cell.width = Twips(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            set_width(tc_pr, "w:tcW", width)
            tc_mar = ensure_child(tc_pr, "w:tcMar")
            for side in ("top", "bottom", "start", "end"):
                node = ensure_child(tc_mar, f"w:{side}")
                node.set(qn("w:w"), str(margins[side]))
                node.set(qn("w:type"), "dxa")


def column_widths_from_weights(weights: list[float], total_width: int) -> list[int]:
    total_weight = sum(weights)
    widths = [round(total_width * weight / total_weight) for weight in weights]
    widths[-1] += total_width - sum(widths)
    return widths


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def style_callout_paragraph(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "5")
    left.set(qn("w:color"), "2E74B5")
    borders.append(left)
    p_pr.append(borders)
    paragraph.paragraph_format.left_indent = Inches(0.15)
    paragraph.paragraph_format.right_indent = Inches(0.15)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(9)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = "1"
    text_run.append(text_node)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph._p.extend([begin, instr, separate, text_run, end])


def configure_section(section, *, landscape=False, wide=False):
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width, section.page_height = Inches(11), Inches(8.5)
    else:
        section.page_width, section.page_height = Inches(8.5), Inches(11)
    margin = Inches(0.5 if wide else 1.0)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = "DigiStayBook | Business & Operational Plan | WIP v3"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in hp.runs:
        set_run_font(run, size=8.5, color=MUTED)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    add_page_field(fp)


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
        "Heading 4": (10.5, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_cover(doc: Document, blocks: list[Block]) -> int:
    first = blocks[0]
    title = str(first.value)
    metadata: list[str] = []
    start = 1
    while start < len(blocks) and blocks[start].kind == "paragraph":
        metadata.append(str(blocks[start].value))
        start += 1
    if start < len(blocks) and blocks[start].kind == "rule":
        start += 1

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(92)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("BUSINESS & OPERATIONAL PLAN")
    set_run_font(run, size=10, color=BLUE, bold=True)
    kicker.paragraph_format.space_after = Pt(16)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    set_run_font(title_run, size=29, color=NAVY, bold=True)
    title_p.paragraph_format.space_after = Pt(14)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Consolidated working plan with stable requirement references")
    set_run_font(sr, size=13, color=MUTED, italic=True)
    sub.paragraph_format.space_after = Pt(42)
    for item in metadata:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline_runs(p, item, size=10)
        p.paragraph_format.space_after = Pt(3)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run("Working draft - legal and tax review gates remain in force")
    set_run_font(nr, size=9, color=MUTED, italic=True)
    note.paragraph_format.space_before = Pt(28)
    doc.add_page_break()

    toc = doc.add_paragraph("Contents", style="Heading 1")
    toc.paragraph_format.space_before = Pt(0)
    for block in blocks:
        if block.kind == "heading" and block.level == 1 and block is not first:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(5)
            run = p.add_run(str(block.value))
            set_run_font(run, size=10.5, color=DARK_BLUE, bold=True)
    doc.add_page_break()
    return start


def heading_paragraph(doc: Document, text: str, level: int):
    label, requirement = split_requirement(text)
    style = f"Heading {min(level, 4)}"
    p = doc.add_paragraph(style=style)
    main = p.add_run(label)
    main.bold = True
    if requirement:
        badge = p.add_run(f"  {requirement}")
        set_run_font(badge, name="Consolas", size=8, color=MUTED, bold=False)
    return p


def table_weights(rows: list[list[str]]) -> list[float]:
    cols = len(rows[0])
    lengths = []
    for index in range(cols):
        maximum = max(len(row[index]) if index < len(row) else 0 for row in rows)
        lengths.append(max(1.0, min(4.5, math.sqrt(maximum + 12) / 3.0)))
    return lengths


def format_table(doc: Document, rows: list[list[str]], total_width: int, font_size: float):
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        for col_index in range(cols):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.1
            add_inline_runs(p, values[col_index] if col_index < len(values) else "", size=font_size)
            if row_index == 0:
                set_cell_fill(cell, LIGHT_BLUE)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = NAVY
    set_repeat_table_header(table.rows[0])
    widths = column_widths_from_weights(table_weights(rows), total_width)
    apply_table_geometry(table, widths, total_width, 120, {"top": 90, "bottom": 90, "start": 120, "end": 120})
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)


def add_record_cards(doc: Document, rows: list[list[str]]):
    headers = rows[0]
    note = doc.add_paragraph()
    note_run = note.add_run("Readable DOCX view: the wide source matrix is presented as one record per data class; all source columns are preserved.")
    set_run_font(note_run, size=9, color=MUTED, italic=True)
    for values in rows[1:]:
        heading_paragraph(doc, values[0], 3)
        for index, label in enumerate(headers[1:], start=1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            p.paragraph_format.space_after = Pt(3)
            label_run = p.add_run(f"{label}: ")
            set_run_font(label_run, size=9, color=NAVY, bold=True)
            add_inline_runs(p, values[index] if index < len(values) else "", size=9.5)
        divider = doc.add_paragraph()
        divider.paragraph_format.space_after = Pt(3)


def build_docx(blocks: list[Block], output: Path):
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    start = add_cover(doc, blocks)
    landscape_active = False

    for block in blocks[start:]:
        if block.kind == "heading":
            heading_paragraph(doc, str(block.value), block.level)
        elif block.kind == "paragraph":
            p = doc.add_paragraph()
            add_inline_runs(p, str(block.value))
        elif block.kind in ("bullets", "numbers"):
            style = "List Bullet" if block.kind == "bullets" else "List Number"
            for item in block.value:
                p = doc.add_paragraph(style=style)
                add_inline_runs(p, str(item))
        elif block.kind == "quote":
            p = doc.add_paragraph()
            style_callout_paragraph(p)
            add_inline_runs(p, str(block.value), size=10)
            for run in p.runs:
                run.italic = True
        elif block.kind == "code":
            lang = block.value["lang"]
            code = block.value["code"]
            p = doc.add_paragraph()
            run = p.add_run(f"[{lang} code block]\n{code}")
            run.font.name = "Courier New"
        elif block.kind == "rule":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
        elif block.kind == "table":
            rows = block.value
            cols = len(rows[0])
            if cols >= 6:
                add_record_cards(doc, rows)
            elif cols >= 4:
                section = doc.add_section(WD_SECTION.NEW_PAGE)
                configure_section(section, landscape=True, wide=True)
                landscape_active = True
                format_table(doc, rows, 14400, 8.5)
                section = doc.add_section(WD_SECTION.NEW_PAGE)
                configure_section(section)
                landscape_active = False
            else:
                format_table(doc, rows, 9360, 9)

    if landscape_active:
        configure_section(doc.add_section(WD_SECTION.NEW_PAGE))
    core = doc.core_properties
    core.title = "DigiStayBook Consolidated Business & Operational Plan v3"
    core.subject = "Authoritative working BOP with stable requirement IDs"
    core.author = "DigiStayBook"
    core.keywords = "DigiStayBook, BOP, requirements, operations, product"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def inline_html(text: str) -> str:
    escaped = html.escape(text)
    escaped = re.sub(r"`([^`]+)`", r"<code>\1</code>", escaped)
    escaped = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", escaped)
    escaped = re.sub(r"\[COLOR:([^\]]+)\](.*?)\[/COLOR\]", r'<span style="color: \1">\2</span>', escaped)
    escaped = re.sub(r"&lt;u&gt;(.*?)&lt;/u&gt;", r"<u>\1</u>", escaped)
    escaped = re.sub(r"\[COLOR:([^\]]+)\](.*?)\[/COLOR\]", r'<span style="color: \1">\2</span>', escaped)
    escaped = re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r"<em>\1</em>", escaped)
    return escaped


def slugify(text: str) -> str:
    label, requirement = split_requirement(text)
    if requirement:
        return requirement
    slug = re.sub(r"[^a-z0-9]+", "-", label.lower()).strip("-")
    return slug or "section"


def blocks_to_html(blocks: list[Block]) -> tuple[str, str]:
    body: list[str] = []
    toc: list[str] = []
    for block in blocks:
        if block.kind == "heading":
            label, requirement = split_requirement(str(block.value))
            anchor = slugify(str(block.value))
            badge = f'<span class="req-id">{html.escape(requirement)}</span>' if requirement else ""
            body.append(f'<h{block.level} id="{anchor}">{inline_html(label)} {badge}</h{block.level}>')
            if block.level in (1, 2):
                toc.append(
                    f'<a class="toc-l{block.level}" href="#{anchor}">{html.escape(label)}'
                    + (f'<small>{html.escape(requirement)}</small>' if requirement else "")
                    + "</a>"
                )
        elif block.kind == "paragraph":
            body.append(f"<p>{inline_html(str(block.value))}</p>")
        elif block.kind in ("bullets", "numbers"):
            tag = "ul" if block.kind == "bullets" else "ol"
            items = "".join(f"<li>{inline_html(str(item))}</li>" for item in block.value)
            body.append(f"<{tag}>{items}</{tag}>")
        elif block.kind == "quote":
            body.append(f"<blockquote>{inline_html(str(block.value))}</blockquote>")
        elif block.kind == "code":
            lang = block.value["lang"]
            code = block.value["code"]
            if lang == "mermaid":
                body.append(f'<pre class="mermaid">{html.escape(code)}</pre>')
            else:
                body.append(f'<pre><code class="language-{html.escape(lang)}">{html.escape(code)}</code></pre>')
        elif block.kind == "rule":
            body.append("<hr>")
        elif block.kind == "table":
            rows = block.value
            header_cells = "".join(f"<th>{inline_html(cell)}</th>" for cell in rows[0])
            data_rows = "".join(
                "<tr>" + "".join(f"<td>{inline_html(cell)}</td>" for cell in row) + "</tr>"
                for row in rows[1:]
            )
            body.append(f'<div class="table-wrap"><table><thead><tr>{header_cells}</tr></thead><tbody>{data_rows}</tbody></table></div>')
    return "\n".join(body), "\n".join(toc)


BASE_CSS = r"""
:root{--navy:#0b2545;--blue:#2e74b5;--ink:#172033;--muted:#5c6778;--line:#d9e1ea;--paper:#fff;--wash:#f5f7fa;--add:#e8f7ed;--del:#fdebec}
*{box-sizing:border-box}html{scroll-behavior:smooth}body{margin:0;background:var(--wash);color:var(--ink);font:16px/1.62 system-ui,-apple-system,"Segoe UI",sans-serif}
a{color:var(--blue)}.layout{display:grid;grid-template-columns:310px minmax(0,1fr);min-height:100vh}.sidebar{position:sticky;top:0;height:100vh;overflow:auto;padding:24px 20px;background:var(--navy);color:#fff}.brand{font-size:20px;font-weight:750;margin-bottom:4px}.side-note{font-size:12px;color:#c9d7e8;margin-bottom:18px}.sidebar a{display:block;color:#eaf2fa;text-decoration:none;border-left:2px solid transparent;padding:6px 8px}.sidebar a:hover{background:#153657;border-left-color:#8fc3ef}.toc-l1{font-weight:700;margin-top:8px}.toc-l2{font-size:13px;padding-left:18px!important}.sidebar small{display:block;color:#9fc1dc;font:10px/1.3 ui-monospace,monospace}.content{width:min(1120px,calc(100% - 48px));margin:28px auto 80px;background:var(--paper);padding:54px 64px;border:1px solid var(--line);box-shadow:0 10px 28px rgba(20,38,60,.08)}
h1,h2,h3,h4{color:var(--navy);line-height:1.2;scroll-margin-top:20px}h1{font-size:30px;margin:52px 0 18px;border-bottom:2px solid var(--line);padding-bottom:9px}h1:first-child{font-size:42px;margin-top:0;border:0}h2{font-size:23px;margin:38px 0 13px}h3{font-size:18px;margin:28px 0 10px}h4{font-size:16px;margin:22px 0 8px}.req-id{display:inline-block;vertical-align:middle;margin-left:8px;padding:3px 7px;border-radius:999px;background:#e8eef5;color:#315d82;font:11px/1.2 ui-monospace,SFMono-Regular,Consolas,monospace;white-space:nowrap}p{margin:0 0 13px}ul,ol{padding-left:25px}li{margin:6px 0}code{font:13px ui-monospace,SFMono-Regular,Consolas,monospace;background:#edf1f5;padding:2px 5px;border-radius:4px}blockquote{margin:18px 0;padding:14px 18px;border-left:4px solid var(--blue);background:#f2f6fa;color:#30445c}.table-wrap{overflow:auto;margin:20px 0 30px;border:1px solid var(--line);border-radius:7px}table{border-collapse:collapse;width:100%;font-size:13px;line-height:1.45}th,td{padding:10px 12px;text-align:left;vertical-align:top;border:1px solid var(--line);min-width:145px}th{position:sticky;top:0;background:#e8eef5;color:var(--navy);font-weight:750}hr{border:0;border-top:1px solid var(--line);margin:24px 0}.top-actions{display:flex;gap:10px;flex-wrap:wrap;margin-bottom:28px}.button{display:inline-block;padding:8px 12px;border-radius:6px;background:var(--navy);color:#fff;text-decoration:none;font-weight:650;font-size:13px}
@media(max-width:900px){.layout{display:block}.sidebar{position:relative;height:auto}.toc-l2{display:none!important}.content{width:calc(100% - 24px);margin:12px;padding:30px 24px}h1:first-child{font-size:34px}}
@media print{body{background:#fff}.layout{display:block}.sidebar,.top-actions{display:none}.content{width:auto;margin:0;padding:0;border:0;box-shadow:none}.table-wrap{overflow:visible}h1{break-before:page}}
"""


def html_shell(
    title: str,
    body: str,
    toc: str,
    extra_css: str = "",
    *,
    side_note: str = "BOP v3 readable view",
    actions: str | None = None,
) -> str:
    if actions is None:
        actions = '<a class="button" href="digistaybook_WIP_v2-to-v3-diff.html">View v2 → v3 changes</a><a class="button" href="digistaybook_WIP_v3.docx">Open DOCX</a>'
    return f"""<!doctype html>
<html lang="en"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>{html.escape(title)}</title><style>{BASE_CSS}{extra_css}</style>\n<script type="module">import mermaid from 'https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.esm.min.mjs'; mermaid.initialize({{ startOnLoad: true }});</script>\n</head>
<body><div class="layout"><nav class="sidebar"><div class="brand">DigiStayBook</div><div class="side-note">{html.escape(side_note)}</div>{toc}</nav>
<main class="content"><div class="top-actions">{actions}</div>{body}</main></div></body></html>"""


@dataclass(frozen=True)
class DiffLine:
    text: str
    source_line: int


def prepare_diff_lines(text: str) -> list[DiffLine]:
    lines: list[DiffLine] = []
    for source_line, line in enumerate(text.splitlines(), start=1):
        if line.startswith("**Stable requirement ID convention:**"):
            continue
        cleaned = REQ_RE.sub("", line).rstrip()
        if not cleaned and lines and not lines[-1].text:
            continue
        lines.append(DiffLine(cleaned, source_line))
    return lines


def nearest_heading(lines: list[DiffLine], index: int) -> str:
    for pos in range(min(index, len(lines) - 1), -1, -1):
        match = HEADING_RE.match(lines[pos].text)
        if match:
            return match.group(2)
    return "Document opening"


def source_range(lines: list[DiffLine], start: int, end: int) -> str:
    if start < end:
        return f"{lines[start].source_line}-{lines[end - 1].source_line}"
    if start < len(lines):
        anchor = lines[start].source_line
    elif lines:
        anchor = lines[-1].source_line + 1
    else:
        anchor = 1
    return f"{anchor}-{anchor}"


def build_diff_html(v2_text: str, v3_text: str) -> str:
    old = prepare_diff_lines(v2_text)
    new = prepare_diff_lines(v3_text)
    matcher = difflib.SequenceMatcher(
        a=[line.text for line in old],
        b=[line.text for line in new],
        autojunk=False,
    )
    changes: list[str] = []
    added = deleted = replaced = 0
    change_number = 0
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        change_number += 1
        if tag == "insert":
            added += j2 - j1
        elif tag == "delete":
            deleted += i2 - i1
        else:
            replaced += 1
            added += j2 - j1
            deleted += i2 - i1
        heading = nearest_heading(new if j1 < len(new) else old, j1 if j1 < len(new) else i1)
        before = "\n".join(line.text for line in old[i1:i2]) or "(no text)"
        after = "\n".join(line.text for line in new[j1:j2]) or "(no text)"
        changes.append(
            f'<section class="change" id="change-{change_number}"><h2>Change {change_number}: {html.escape(heading)}</h2>'
            f'<div class="line-ref">v2 lines {source_range(old, i1, i2)} → v3 lines {source_range(new, j1, j2)}</div>'
            f'<div class="compare"><div class="before"><h3>v2</h3><pre>{html.escape(before)}</pre></div>'
            f'<div class="after"><h3>v3</h3><pre>{html.escape(after)}</pre></div></div></section>'
        )
    nav = "".join(f'<a href="#change-{i}">Change {i}</a>' for i in range(1, change_number + 1))
    summary = f"""
<h1>DigiStayBook BOP: v2 → v3 content changes</h1>
<p>This standalone comparison shows wording and product-content changes only. Requirement-ID suffixes and the ID-convention metadata are intentionally excluded, and unchanged content is hidden.</p>
<div class="stats"><div><strong>{change_number}</strong><span>change blocks</span></div><div><strong>{added}</strong><span>v3 lines added</span></div><div><strong>{deleted}</strong><span>v2 lines removed</span></div><div><strong>{replaced}</strong><span>replacement blocks</span></div></div>
<p><a class="button" href="digistaybook_WIP_v3.html">Read the complete v3 plan</a></p>
"""
    css = r"""
.layout{grid-template-columns:230px minmax(0,1fr)}.content{width:min(1440px,calc(100% - 36px));padding:42px}.sidebar a{font-size:12px}.stats{display:grid;grid-template-columns:repeat(4,minmax(120px,1fr));gap:12px;margin:22px 0 32px}.stats div{background:#eef3f8;border:1px solid var(--line);border-radius:8px;padding:15px}.stats strong{display:block;font-size:26px;color:var(--navy)}.stats span{font-size:12px;color:var(--muted)}.change{border-top:2px solid var(--line);padding-top:16px;margin-top:38px}.change h2{margin-top:0}.line-ref{font:12px ui-monospace,monospace;color:var(--muted);margin-bottom:10px}.compare{display:grid;grid-template-columns:1fr 1fr;gap:14px}.compare>div{min-width:0;border-radius:7px;border:1px solid var(--line);overflow:hidden}.compare h3{margin:0;padding:9px 12px;font-size:14px}.before h3{background:#f8dfe1}.after h3{background:#dff2e5}pre{white-space:pre-wrap;overflow-wrap:anywhere;margin:0;padding:14px;font:12px/1.5 ui-monospace,SFMono-Regular,Consolas,monospace;max-height:520px;overflow:auto}.before pre{background:var(--del)}.after pre{background:var(--add)}
@media(max-width:900px){.compare{grid-template-columns:1fr}.stats{grid-template-columns:1fr 1fr}}
"""
    actions = '<a class="button" href="digistaybook_WIP_v3.html">Read complete v3</a><a class="button" href="digistaybook_WIP_v3.docx">Open DOCX</a>'
    return html_shell(
        "DigiStayBook BOP v2 to v3 content comparison",
        summary + "\n".join(changes),
        nav,
        css,
        side_note="v2 to v3 content changes",
        actions=actions,
    )


def validate_ids(blocks: list[Block]):
    ids: list[str] = []
    missing: list[str] = []
    for block in blocks:
        if block.kind == "heading" and 2 <= block.level <= 4:
            _, requirement = split_requirement(str(block.value))
            if requirement:
                ids.append(requirement)
            else:
                missing.append(str(block.value))
    duplicates = sorted({item for item in ids if ids.count(item) > 1})
    if missing or duplicates:
        raise SystemExit(f"Requirement ID validation failed. Missing={missing}; duplicates={duplicates}")
    return len(ids)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--v2", type=Path, default=Path("digistaybook_WIP_v2.md"))
    parser.add_argument("--v3", type=Path, default=Path("digistaybook_WIP_v3.md"))
    parser.add_argument("--outdir", type=Path, default=Path("artifacts"))
    parser.add_argument("--diff-only", action="store_true", help="Rebuild only the content-only v2-to-v3 comparison")
    args = parser.parse_args()

    v2_text = args.v2.read_text(encoding="utf-8")
    v3_text = args.v3.read_text(encoding="utf-8")
    blocks = parse_markdown(v3_text)
    id_count = validate_ids(blocks)
    args.outdir.mkdir(parents=True, exist_ok=True)

    docx_path = args.outdir / "digistaybook_WIP_v3.docx"
    html_path = args.outdir / "digistaybook_WIP_v3.html"
    diff_path = args.outdir / "digistaybook_WIP_v2-to-v3-diff.html"
    if not args.diff_only:
        build_docx(blocks, docx_path)
        body, toc = blocks_to_html(blocks)
        html_path.write_text(html_shell("DigiStayBook BOP v3", body, toc), encoding="utf-8")
    diff_path.write_text(build_diff_html(v2_text, v3_text), encoding="utf-8")
    if not args.diff_only:
        print(f"Built {docx_path}")
        print(f"Built {html_path}")
    print(f"Built {diff_path}")
    print(f"Validated {id_count} unique requirement/template IDs")


if __name__ == "__main__":
    main()
